from pathlib import Path

from docx import Document


def create_word_document(text: str, filename: str = "sample.docx", save_dir: Path | None = None) -> Path:
    """Create or update a Word document and save it in the Downloads folder."""
    target_dir = save_dir or (Path.home() / "Downloads")
    target_dir.mkdir(parents=True, exist_ok=True)

    file_path = target_dir / filename
    if file_path.exists():
        document = Document(file_path)
    else:
        document = Document()

    document.add_paragraph(text)
    document.save(file_path)
    return file_path


if __name__ == "__main__":
    output_path = create_word_document("This is a test written from Python.")
    print(f"Word file saved to: {output_path}")
